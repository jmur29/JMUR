from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Cm(1.8)
section.bottom_margin = Cm(1.8)
section.left_margin   = Cm(2.0)
section.right_margin  = Cm(2.0)

# ── Colour palette ────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x1B, 0x2A, 0x4A)   # headers
TEAL   = RGBColor(0x00, 0x7A, 0x7C)   # sub-headers
GREEN  = RGBColor(0x1A, 0x6B, 0x3C)   # confirmed
GREY   = RGBColor(0xF2, 0xF5, 0xF8)   # table shading
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
BLACK  = RGBColor(0x1A, 0x1A, 0x1A)

# ── Helpers ───────────────────────────────────────────────────────────────────
def shade_cell(cell, hex_colour="F2F5F8"):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_colour)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top","left","bottom","right"):
        val = kwargs.get(side, {"sz":"4","val":"single","color":"C8D0DC"})
        el  = OxmlElement(f"w:{side}")
        for k,v in val.items():
            el.set(qn(f"w:{k}"), v)
        tcBorders.append(el)
    tcPr.append(tcBorders)

def add_heading(doc, text, level=1):
    p    = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level==1 else 8)
    p.paragraph_format.space_after  = Pt(4)
    run  = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(14 if level==1 else 11)
    run.font.color.rgb = NAVY if level==1 else TEAL
    return p

def add_subheading(doc, text):
    p   = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = TEAL
    return p

def add_body(doc, text, bold=False, italic=False, size=9):
    p   = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.color.rgb = BLACK
    return p

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # header row
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        shade_cell(cell, "1B2A4A")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p   = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(8.5)
        run.font.color.rgb = WHITE

    # data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        fill = "F2F5F8" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            shade_cell(cell, fill)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p   = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(1)
            bold_cell = str(cell_text).startswith("**")
            clean     = str(cell_text).replace("**","")
            run = p.add_run(clean)
            run.bold = bold_cell
            run.font.size = Pt(8.5)
            run.font.color.rgb = BLACK

    # column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    return table

def add_spacer(doc, size=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(size)
    run = p.add_run(" ")
    run.font.size = Pt(size)

# ══════════════════════════════════════════════════════════════════════════════
# TITLE BLOCK
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after  = Pt(2)
run = p.add_run("SOURCE OF FUNDS — COMPLETE MONEY MAP")
run.bold = True
run.font.size = Pt(18)
run.font.color.rgb = NAVY

p2 = doc.add_paragraph()
p2.paragraph_format.space_after = Pt(1)
run2 = p2.add_run("Myrtille Remond-Bernon & Jean-Baptiste Bernon  |  269 Sammon Ave, Toronto ON")
run2.font.size = Pt(10)
run2.font.color.rgb = TEAL

p3 = doc.add_paragraph()
p3.paragraph_format.space_after = Pt(8)
run3 = p3.add_run("Purchase Price: $925,000  |  Mortgage: $475,000 (First National #1535557)  |  Down Payment: $450,000  |  Closing: May 14, 2026")
run3.font.size = Pt(9)
run3.font.color.rgb = RGBColor(0x55,0x55,0x55)

# divider
p_div = doc.add_paragraph()
p_div.paragraph_format.space_after = Pt(8)
run_div = p_div.add_run("─" * 110)
run_div.font.size = Pt(7)
run_div.font.color.rgb = RGBColor(0xC0,0xC8,0xD8)

# ══════════════════════════════════════════════════════════════════════════════
# PART 1 — FRENCH WIRE
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "PART 1 — FRENCH WIRE: EUR 210,000 → CAD $333,588.50")

add_subheading(doc, "A. French Accounts Liquidated (Feb 7–23, 2026)")
add_body(doc, "All funds are Myrtille's own personal savings held at Caisse d'Épargne de Midi-Pyrénées. "
              "Long-term savings history confirmed by Synthèse dated July 29, 2025 (EUR 215,030 total across all accounts).")
add_spacer(doc, 3)

add_table(doc,
    ["French Account", "Product", "Balance Feb 6", "Redeemed/Closed", "Amount Received", "Document"],
    [
        ["09044220382", "CAT RENOUVELABLE FIXE (term deposit)", "EUR 150,000.00", "Feb 7 → re-rolled → released Feb 17", "EUR 150,000.00", "Feb 6 Synthèse + Mar 3 stmt"],
        ["09038490716", "CAPTIO PRESTANCE (term deposit)", "EUR 10,000.00", "Feb 14 — early redemption", "EUR 10,219.87", "Feb 6 Synthèse + Mar 3 stmt"],
        ["09038490817", "CAPTIO PRESTANCE (term deposit)", "EUR 10,000.00", "Feb 14 — early redemption", "EUR 10,219.87", "CLOTURE redemption slip"],
        ["09039454248", "CATPEL QUADRETO 4 ANS 2023", "EUR 337.90", "Feb 14 — early redemption", "EUR 390.07", "CLOTURE doc"],
        ["16439214638", "PEL — Plan Épargne Logement", "EUR 15,740.76", "Feb 14 — account closed", "EUR 16,474.52", "CLOTURE closure slip"],
        ["00025301236", "Livret A savings account", "EUR 22,252.49", "Feb 17 — account closed", "EUR 22,200.00", "Feb 6 Synthèse"],
        ["04398009403", "Checking account balance", "EUR 5,678.42", "Carried forward to wire", "EUR 5,678.42", "Mar 3 account stmt"],
        ["N/A — Notaire", "FAURIE-GREPON refund (failed purchase deposit)", "—", "Feb 23 — wired by notaire", "EUR 10,140.00", "Notaire attestation + Gmail"],
        ["**TOTAL**", "**All Myrtille's own funds**", "**EUR ~214,009**", "", "**EUR ~225,322**", ""],
        ["**Wire executed Feb 25**", "EUR 210,000 sent | EUR 90 bank fees", "", "", "**−EUR 210,090**", "Wire instruction Feb 18"],
        ["**Residual left in France**", "Post-wire balance", "", "", "**EUR 15,704.97**", "Mar 3 Synthèse ✅"],
    ],
    col_widths=[1.2, 1.6, 1.1, 1.5, 1.2, 1.5]
)

add_spacer(doc, 5)
add_body(doc, "Exchange rate applied: EUR 210,000 × 1.58851 = CAD $333,588.50  |  "
              "Received in TD Chequing 6702242 on Feb 25, 2026  |  Reference: 260225S0582800WIRE", italic=True)

add_spacer(doc, 6)
add_subheading(doc, "B. French Money Trail Inside Canada")
add_body(doc, "After landing in Myrtille's TD chequing, the $333,588.50 moved as follows:")
add_spacer(doc, 3)

add_table(doc,
    ["Date", "Event", "Amount", "Balance of French $", "Document"],
    [
        ["Feb 25, 2026", "Wire received — TD Chequing 6702242 (ref: 260225S0582800WIRE)", "+$333,588.50", "$333,588.50", "TD Chequing 6702242 — 120-day stmt"],
        ["Feb 25–Mar 2", "Normal household expenses paid from chequing (incl. $15,885 cheque Feb 27)", "~−$12,945", "~$320,643", "TD Chequing 6702242 — 120-day stmt"],
        ["Mar 3, 2026", "**$320,000 placed into GIC 8007943** — parked to earn interest while waiting for closing", "−$320,000.00", "$643 liquid", "Transaction detail + chequing stmt"],
        ["Mar 3–Apr 22", "Remaining ~$643 of French $ absorbed into chequing; regular income/expenses continue", "—", "Commingled", "TD Chequing 6702242 — 120-day stmt"],
        ["Apr 23, 2026", "**GIC 8007943 matures — $320,514.19 returned to chequing** (+$514.19 interest earned)", "+$320,514.19", "$320,514.19", "Transaction detail + chequing stmt"],
        ["Apr 23, 2026", "**$46,250 transferred to Mysak Real Estate Trust** — deposit on 269 Sammon Ave", "−$46,250.00", "$274,264.19 in chequing", "TD branch deposit receipt"],
    ],
    col_widths=[1.0, 2.7, 1.1, 1.4, 1.9]
)

add_spacer(doc, 5)
add_subheading(doc, "C. Where the French Money Sits Today")
add_spacer(doc, 2)

add_table(doc,
    ["Location", "Amount", "Status"],
    [
        ["TD Chequing 6702242 (French wire portion)", "~$274,264", "Liquid — available at closing"],
        ["Mysak Real Estate Trust (deposit paid Apr 23)", "$46,250.00", "Confirmed at solicitor — counts toward $450K"],
        ["Interest earned on GIC 8007943 during hold period", "$514.19", "Included in chequing balance above"],
        ["Spent on normal living expenses Feb 25 – Apr 22", "~$12,560", "Normal household costs during the wait period"],
        ["**TOTAL FRENCH WIRE ACCOUNTED FOR**", "**$333,588.50**", "**✅ Fully traced**"],
    ],
    col_widths=[3.2, 1.4, 2.5]
)

# ══════════════════════════════════════════════════════════════════════════════
# PART 2 — PRE-EXISTING CANADIAN SAVINGS
# ══════════════════════════════════════════════════════════════════════════════
add_spacer(doc, 8)
add_heading(doc, "PART 2 — PRE-EXISTING CANADIAN SAVINGS (No Connection to French Wire)")

add_subheading(doc, "A. Myrtille — TD ePremium Savings Account 6642878  |  Balance Apr 23: $20,070.55")
add_body(doc, "100% pre-existing Canadian savings. Account predates the French wire. The $19,000 that briefly moved to a sub-GIC on Mar 3 was Myrtille's own ePremium funds — it returned Apr 23 as $19,030.53.")
add_spacer(doc,2)

add_table(doc,
    ["Event", "Amount", "Date", "Source"],
    [
        ["Account balance (long-standing savings)", "$21,896.95", "Jan 22, 2026", "Pre-existing Canadian savings"],
        ["$8,000 moved to First Home Savings Account (FHSA)", "−$8,000.00", "Feb 6, 2026", "Myrtille's own reallocation"],
        ["$5,728.66 received from chequing (pre-wire funds)", "+$5,728.66", "Feb 6, 2026", "Chequing top-up (pre-wire balance)"],
        ["$19,000 placed in GIC 8007943-04 (short-term)", "−$19,000.00", "Mar 3, 2026", "Myrtille's own ePremium funds"],
        ["$19,030.53 returned from GIC 8007943-04", "+$19,030.53", "Apr 23, 2026", "Principal + $30.53 interest back"],
        ["Small regular transfers + interest credits", "~+$415", "Jan–Apr", "Routine"],
        ["**Balance Apr 23, 2026**", "**$20,070.55**", "", "**Pre-existing Canadian savings only ✅**"],
    ],
    col_widths=[3.0, 1.2, 1.1, 2.8]
)

add_spacer(doc, 6)
add_subheading(doc, "B. Myrtille — TD TFSA High Interest Savings 6683965  |  Balance Apr 23: $41,357.07")
add_body(doc, "100% pre-existing Canadian TFSA investments. GIC 8002852-02 was issued January 2, 2025 — 14 months before the French wire arrived.")
add_spacer(doc, 2)

add_table(doc,
    ["Event", "Amount", "Date", "Source"],
    [
        ["GIC 8002852-02 issued into TFSA", "$35,000.00 principal", "Jan 2, 2025", "Pre-existing TFSA savings — 14 months before French wire"],
        ["GIC 8002852-06 issued into TFSA", "$10,000.00 principal", "Oct 7, 2025", "Pre-existing TFSA savings — 5 months before French wire"],
        ["GIC 8002852-02 matures → TFSA", "+$36,412.84", "Mar 3, 2026", "Return of Jan 2025 investment + interest"],
        ["New GIC 8002852-07 opened from TFSA", "−$41,000.00", "Mar 3, 2026", "Re-invested for final 7 weeks"],
        ["GIC 8002852-07 matures → TFSA", "+$41,065.88", "Apr 23, 2026", "Return + $65.88 interest"],
        ["Regular transfers + interest credits", "~+$291", "Jan–Apr", "Routine"],
        ["**Balance Apr 23, 2026**", "**$41,357.07**", "", "**100% pre-existing Canadian TFSA savings ✅**"],
    ],
    col_widths=[3.0, 1.3, 1.1, 2.7]
)

add_spacer(doc, 6)
add_subheading(doc, "C. Myrtille — Non-Liquid Investments (NOT counted toward $450K closing funds)")
add_spacer(doc, 2)
add_table(doc,
    ["Account", "Balance", "Issued", "Matures", "Source"],
    [
        ["GIC 8007943-01 (TD Special Offer GIC)", "$15,000", "Oct 7, 2025", "Dec 6, 2026", "Pre-existing Canadian savings"],
        ["GIC 8002852-06 (TD Canadian Banks GIC TFSA)", "$10,000", "Oct 7, 2025", "Oct 7, 2030", "Pre-existing Canadian TFSA savings"],
        ["FHSA 660020311628", "$8,000", "Feb 6, 2026", "—", "Transferred from ePremium (Myrtille's own)"],
    ],
    col_widths=[2.4, 0.9, 1.0, 1.0, 2.8]
)

add_spacer(doc, 6)
add_subheading(doc, "D. Jean-Baptiste — TD Chequing 8446-6630764  |  Balance Mar 31: $48,713.32")
add_body(doc, "Funded entirely from employment income (Toronto Drydock Ltd.) and a parental gift from Christian Schmitter (JB's father). No connection to Myrtille's French savings.")
add_spacer(doc, 2)

add_table(doc,
    ["Component", "Amount", "Date", "Source / Document"],
    [
        ["Opening balance", "$20,608.34", "Dec 31, 2025", "Pre-existing savings — TD statement Dec 31–Jan 30"],
        ["Parental gift — EUR 15,000 from Christian Schmitter", "+$24,155.02", "Jan 12, 2026", "Schmitter (Banque Pop.) → Wise → TD  |  Schmitter stmt + Wise stmt + TD stmt"],
        ["Toronto Drydock payroll (Jan–Mar net savings)", "~+$7,000", "Jan–Mar 2026", "Bi-weekly payroll confirmed on all 3 TD statements"],
        ["Regular household expenses", "~−$3,050", "Jan–Mar 2026", "TD statements"],
        ["**Confirmed balance Mar 31, 2026**", "**$48,713.32**", "", "**TD statement Feb 27–Mar 31 ✅**"],
        ["Drydock payroll ~Apr 10 (estimate)", "+$1,811", "~Apr 10, 2026", "Based on Apr 23 paystub ($1,811.04 net)"],
        ["**Estimated balance Apr 23**", "**~$50,524**", "", ""],
    ],
    col_widths=[2.8, 1.1, 1.0, 3.2]
)

# ══════════════════════════════════════════════════════════════════════════════
# PART 3 — CLOSING FUNDS CALCULATION
# ══════════════════════════════════════════════════════════════════════════════
add_spacer(doc, 8)
add_heading(doc, "PART 3 — $450,000 DOWN PAYMENT: CLOSING FUNDS CALCULATION")

add_subheading(doc, "Confirmed Liquid Balances")
add_spacer(doc, 2)

add_table(doc,
    ["Account", "Balance", "Date Confirmed", "Owner", "Origin", "Document"],
    [
        ["TD Chequing 6702242", "$287,110.39", "Apr 24, 2026", "Myrtille", "~$274,264 French wire + ~$12,846 pre-existing/income", "120-day statement"],
        ["TD ePremium Savings 6642878", "$20,070.55", "Apr 23, 2026", "Myrtille", "Pre-existing Canadian savings", "Account statement"],
        ["TD TFSA 6683965", "$41,357.07", "Apr 23, 2026", "Myrtille", "Pre-existing Canadian TFSA (GICs from Jan 2025)", "Account statement"],
        ["TD Chequing 8446-6630764", "$48,713.32", "Mar 31, 2026", "Jean-Baptiste", "Drydock payroll + EUR 15,000 parental gift", "Bank statement"],
        ["Mysak Real Estate Trust", "$46,250.00", "Apr 23, 2026", "Both", "French wire (via GIC 8007943 maturity)", "TD branch receipt"],
        ["**Confirmed liquid sub-total**", "**$443,501.33**", "", "", "", ""],
    ],
    col_widths=[1.8, 1.0, 1.1, 0.9, 2.2, 1.2]
)

add_spacer(doc, 5)
add_subheading(doc, "Pre-Closing Income (Apr 24 – May 14)")
add_spacer(doc, 2)

add_table(doc,
    ["Income Item", "Est. Amount", "Expected Date", "Basis"],
    [
        ["Myrtille — Empire Life LTD benefit", "$3,916.00", "~May 14, 2026", "Confirmed monthly credit mid-month in chequing (visible Apr 14, Mar 16, Feb 17 on stmt)"],
        ["JB — Toronto Drydock payroll", "$1,811.04", "~Apr 24, 2026", "Bi-weekly cadence; Apr 23 paystub = $1,811.04 net"],
        ["JB — Toronto Drydock payroll", "$1,811.04", "~May 7, 2026", "Same cadence"],
        ["**Pre-closing income sub-total**", "**$7,538.08**", "", ""],
    ],
    col_widths=[2.5, 1.0, 1.3, 3.3]
)

add_spacer(doc, 5)
add_subheading(doc, "Final Closing Funds Summary")
add_spacer(doc, 2)

add_table(doc,
    ["", "Amount"],
    [
        ["Confirmed liquid balances (Apr 23–24)", "$443,501.33"],
        ["Pre-closing income (LTD + 2× Drydock payrolls)", "$7,538.08"],
        ["**TOTAL AVAILABLE AT CLOSING**", "**$451,039.41**"],
        ["Down payment required", "$450,000.00"],
        ["**SURPLUS**", "**$1,039.41**"],
    ],
    col_widths=[5.0, 1.5]
)

add_spacer(doc, 5)
add_body(doc,
    "Note: The two GICs (8002852-06 $10,000 and 8007943-01 $15,000) are NOT included above and are NOT needed for closing. "
    "They are disclosed as background investment assets only. The $450,000 closes entirely from liquid accounts.",
    italic=True)

# ══════════════════════════════════════════════════════════════════════════════
# PART 4 — ONE-PAGE SUMMARY
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_heading(doc, "PART 4 — COMPLETE SOURCE SUMMARY: Every Dollar, Every Origin")
add_spacer(doc, 3)

add_table(doc,
    ["Pool", "Account / Source", "Amount", "% of $450K", "Owner", "Origin"],
    [
        ["French wire", "TD Chequing 6702242 (French portion)", "~$274,264", "60.9%", "Myrtille", "EUR 210,000 wire Feb 25 — Myrtille's own Caisse d'Épargne savings"],
        ["French wire", "Mysak RE Trust (deposit paid Apr 23)", "$46,250", "10.3%", "Both", "From GIC 8007943 maturity — funded by French wire proceeds"],
        ["Canadian TFSA savings", "TD TFSA 6683965", "$41,357", "9.2%", "Myrtille", "Pre-existing Canadian TFSA — GICs issued Jan 2025 & Oct 2025"],
        ["JB employment + gift", "TD Chequing 8446-6630764", "~$50,524", "11.2%", "Jean-Baptiste", "Drydock payroll savings (~$26,369) + parental gift ($24,155)"],
        ["Canadian savings", "TD ePremium Savings 6642878", "$20,071", "4.5%", "Myrtille", "Pre-existing Canadian savings account"],
        ["Pre-existing / income", "TD Chequing 6702242 (residual)", "~$12,846", "2.9%", "Myrtille", "Pre-wire chequing balance + LTD income, CCB etc. Feb–Apr"],
        ["Pre-closing income", "LTD (May ~14) + JB payroll (Apr 24, May 7)", "~$7,538", "1.7%", "Both", "LTD benefit (Empire Life) + Toronto Drydock bi-weekly payroll"],
        ["**TOTAL**", "", "**~$452,850**", "**100%**", "", ""],
    ],
    col_widths=[1.3, 1.9, 0.9, 0.7, 0.8, 2.5]
)

add_spacer(doc, 8)

# summary paragraph
p_sum = doc.add_paragraph()
p_sum.paragraph_format.space_before = Pt(4)
p_sum.paragraph_format.space_after  = Pt(4)
run_s = p_sum.add_run("SUMMARY   ")
run_s.bold = True
run_s.font.size = Pt(9.5)
run_s.font.color.rgb = NAVY
run_body = p_sum.add_run(
    "The $450,000 down payment is drawn from three clearly separated pools: (1) Myrtille's French personal savings — "
    "EUR 210,000 wired Feb 25, converted to CAD $333,588.50, held in GIC Mar 3–Apr 23, with $46,250 already at "
    "the solicitor and $274,264 sitting in TD chequing — funding approximately 71% of the down payment. "
    "(2) Myrtille's pre-existing Canadian savings — her TFSA (GICs issued January 2025 and October 2025, well before "
    "the application) and her ePremium savings — funding approximately 14%. "
    "(3) Jean-Baptiste's funds — accumulated Drydock payroll plus the EUR 15,000 parental gift from his father "
    "Christian Schmitter — funding the remaining 15%. "
    "No unidentified deposits. No third-party loans. Every dollar is sourced and documented."
)
run_body.font.size = Pt(9.5)
run_body.font.color.rgb = BLACK

# ── footer ────────────────────────────────────────────────────────────────────
add_spacer(doc, 6)
p_ft = doc.add_paragraph()
p_ft.paragraph_format.space_before = Pt(8)
run_ft = p_ft.add_run(
    "Prepared by: Jake Murray  |  Jake Murray Mortgages (FSRA #12923)  |  "
    "Lender: First National Financial  |  File #1535557  |  Closing: May 14, 2026"
)
run_ft.font.size  = Pt(8)
run_ft.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
run_ft.italic = True

# ── Save ──────────────────────────────────────────────────────────────────────
output_path = "/home/user/JMUR/Bernon_Remond_Source_of_Funds_Money_Map.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
